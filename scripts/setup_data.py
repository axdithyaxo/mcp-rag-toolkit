import csv
import os
from pathlib import Path
import pickle
import faiss
from sentence_transformers import SentenceTransformer
from pdfminer.high_level import extract_text
from docx import Document

# Model and output paths
MODEL = SentenceTransformer("all-MiniLM-L6-v2")
DATA_ROOT = Path("data/enterprise_rag_sample_docs_v3")
INDEX_PATH = "index/faiss.index"
DOCS_PATH = "index/doc_texts.pkl"

# Supported file extensions
TEXT_EXTENSIONS = [".txt", ".md", ".html", ".csv", ".pdf", ".docx"]
skipped_reasons = {}

def extract_text_from_file(file_path):
    try:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return extract_text(str(file_path))
        elif suffix == ".docx":
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                print(f"⚠️ Skipped (empty docx): {file_path}")
                return None
            return "\n".join(paragraphs)
        elif suffix == ".docx":
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                    print(f"⚠️ Skipped (empty docx): {file_path}")
                    return None
            full_text = "\n".join(paragraphs)
            print(f"📄 Indexed DOCX: {file_path} ({len(full_text)} chars)")
            return full_text
        else:
            return Path(file_path).read_text(encoding="utf-8")
    except Exception as e:
        skipped_reasons[str(file_path)] = str(e)
        print(f"⚠️ Skipped: {file_path} — {e}")
        return None

def index_documents():
    documents = []
    filepaths = []
    all_files = []

    for file in DATA_ROOT.rglob("*"):
        if file.suffix.lower() in TEXT_EXTENSIONS and file.is_file():
            all_files.append(file)
            content = extract_text_from_file(file)
            if content:
                documents.append(content)
                filepaths.append(str(file))
            else:
                skipped_reasons[str(file)] = "Empty or unreadable"

    # Create embeddings
    embeddings = MODEL.encode(documents, convert_to_tensor=False)

    # Store index
    index = faiss.IndexFlatL2(len(embeddings[0]))
    index.add(embeddings)
    faiss.write_index(index, INDEX_PATH)

    # Save paths
    with open(DOCS_PATH, "wb") as f:
        pickle.dump(filepaths, f)

    # Write indexing log
    log_path = Path("logs/index_log.csv")
    csv_log_scripts_path = Path("scripts/index_log.csv")
    with open(log_path, "w", newline="", encoding="utf-8") as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["file_path", "status", "reason"])
        for i in range(len(all_files)):
            file = all_files[i]
            if str(file) in filepaths:
                writer.writerow([str(file), "Indexed", ""])
            else:
                reason = skipped_reasons.get(str(file), "Unknown")
                writer.writerow([str(file), "Skipped", reason])
    # Duplicate CSV log to scripts/
    with open(csv_log_scripts_path, "w", newline="", encoding="utf-8") as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["file_path", "status", "reason"])
        for i in range(len(all_files)):
            file = all_files[i]
            if str(file) in filepaths:
                writer.writerow([str(file), "Indexed", ""])
            else:
                reason = skipped_reasons.get(str(file), "Unknown")
                writer.writerow([str(file), "Skipped", reason])

    # Write indexing log as plain text
    txt_log_path = Path("logs/index_log.txt")
    txt_log_scripts_path = Path("scripts/index_log.txt")
    with open(txt_log_path, "w", encoding="utf-8") as txt_file:
        for file in all_files:
            if str(file) in filepaths:
                txt_file.write(f"Indexed: {file}\n")
            else:
                reason = skipped_reasons.get(str(file), "Unknown")
                txt_file.write(f"Skipped: {file} — Reason: {reason}\n")
    # Duplicate TXT log to scripts/
    with open(txt_log_scripts_path, "w", encoding="utf-8") as txt_file:
        for file in all_files:
            if str(file) in filepaths:
                txt_file.write(f"Indexed: {file}\n")
            else:
                reason = skipped_reasons.get(str(file), "Unknown")
                txt_file.write(f"Skipped: {file} — Reason: {reason}\n")

    print(f"✅ Indexed {len(filepaths)} documents into FAISS.")

if __name__ == "__main__":
    index_documents()